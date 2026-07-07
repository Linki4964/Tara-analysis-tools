#!/usr/bin/env python3
"""
AI 自动 DOCX → JSON 转换器
自动读取当前文件夹下的所有 .docx 文件，使用 AI 智能解析并输出结构化 JSON。
"""

import os
import json
import sys
from pathlib import Path
from docx import Document
from openai import OpenAI

# ============================================================
# 配置 - 在这里修改你的 API 设置
# ============================================================
API_KEY = "sk-1bb0d4ef94bc4e12a9b718a4dc196915"
API_BASE_URL = "https://api.deepseek.com"   # DeepSeek API
MODEL_NAME = "deepseek-chat"                 # DeepSeek 模型

# 当前脚本所在目录即为工作目录
WORK_DIR = Path(__file__).parent.resolve()
OUTPUT_DIR = WORK_DIR / "output_json"


def extract_docx_content(filepath: Path) -> dict:
    """
    从 DOCX 文件中提取所有内容：
    - 段落（带样式名）
    - 表格（逐行逐列）
    - 文件元信息
    """
    doc = Document(str(filepath))

    # 提取段落
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "style": para.style.name if para.style else "Normal",
                "text": text
            })

    # 提取表格
    tables = []
    for ti, table in enumerate(doc.tables):
        rows_data = []
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append({"row_index": ri, "cells": cells})
        tables.append({
            "table_index": ti,
            "row_count": len(table.rows),
            "col_count": len(table.columns),
            "rows": rows_data
        })

    return {
        "filename": filepath.name,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables
    }


def build_ai_prompt(doc_data: dict) -> str:
    """
    根据提取的文档内容构建 AI 提示词
    """
    parts = [f"文件名: {doc_data['filename']}"]

    if doc_data["paragraphs"]:
        parts.append("\n--- 段落内容 ---")
        for p in doc_data["paragraphs"]:
            parts.append(f"[{p['style']}] {p['text']}")

    if doc_data["tables"]:
        parts.append("\n--- 表格内容 ---")
        for t in doc_data["tables"]:
            parts.append(f"\n表格 {t['table_index'] + 1} ({t['row_count']}行 × {t['col_count']}列):")
            for row in t["rows"]:
                parts.append(" | ".join(row["cells"]))

    raw_text = "\n".join(parts)

    prompt = f"""你是一个专业的数据结构分析师。请将以下从 DOCX 文档中提取的内容转换为结构化的 JSON 格式。

要求:
1. 仔细分析内容，识别出数据的逻辑结构（如：需求列表、配置项、检查清单等）
2. 如果有表格，将表格数据转为 JSON 数组，每行一个对象
3. 保留所有原始文本内容，不要遗漏任何信息
4. 添加有意义的字段名（使用英文 camelCase）
5. 如果有层级关系，用嵌套 JSON 表示
6. 返回纯 JSON，不要包含任何解释文字，不要用 markdown 代码块包裹

文档内容如下:
{raw_text}"""

    return prompt


def ai_convert_to_json(doc_data: dict) -> dict:
    """
    调用 AI API 将文档内容转换为结构化 JSON
    """
    client = OpenAI(
        api_key=API_KEY,
        base_url=API_BASE_URL
    )

    prompt = build_ai_prompt(doc_data)

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "你是一个精确的 JSON 数据转换器。只返回 JSON，不要有其他内容。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=4096
        )

        result_text = response.choices[0].message.content.strip()

        # 清理可能的 markdown 代码块标记
        if result_text.startswith("```"):
            lines = result_text.split("\n")
            # 移除首行 ```json 和末行 ```
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            result_text = "\n".join(lines)

        ai_result = json.loads(result_text)
        return ai_result

    except json.JSONDecodeError as e:
        print(f"  ⚠️  AI 返回的内容不是有效 JSON: {e}")
        print(f"  原始返回: {result_text[:500]}...")
        # 回退：返回原始提取数据
        return {"error": "AI JSON parse failed", "raw": result_text, "extractedData": doc_data}
    except Exception as e:
        print(f"  ⚠️  AI 调用失败: {e}")
        return {"error": str(e), "extractedData": doc_data}


def process_docx(filepath: Path) -> bool:
    """
    处理单个 DOCX 文件：提取 → AI 转换 → 输出 JSON
    """
    print(f"\n{'='*60}")
    print(f"📄 处理文件: {filepath.name}")
    print(f"{'='*60}")

    # 第一步：提取原始内容
    print("  📋 提取文档内容...")
    doc_data = extract_docx_content(filepath)
    print(f"     - {doc_data['paragraph_count']} 个段落")
    print(f"     - {doc_data['table_count']} 个表格")

    # 第二步：保存原始提取（不用 AI，纯本地提取的 JSON）
    raw_output_path = OUTPUT_DIR / f"{filepath.stem}_raw.json"
    raw_output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(raw_output_path, "w", encoding="utf-8") as f:
        json.dump(doc_data, f, ensure_ascii=False, indent=2)
    print(f"  💾 原始提取已保存: {raw_output_path.name}")

    # 第三步：AI 智能转换
    print(f"  🤖 调用 AI 进行智能结构化转换...")
    ai_result = ai_convert_to_json(doc_data)

    # 第四步：保存 AI 转换结果
    ai_output_path = OUTPUT_DIR / f"{filepath.stem}_ai_structured.json"
    with open(ai_output_path, "w", encoding="utf-8") as f:
        json.dump(ai_result, f, ensure_ascii=False, indent=2)
    print(f"  ✅ AI 结构化结果已保存: {ai_output_path.name}")

    return True


def main():
    print("=" * 60)
    print("🤖 AI DOCX → JSON 智能转换器")
    print("=" * 60)
    print(f"工作目录: {WORK_DIR}")
    print(f"API 地址: {API_BASE_URL}")
    print(f"模型: {MODEL_NAME}")

    # 扫描当前目录下的所有 DOCX 文件
    docx_files = sorted(WORK_DIR.glob("*.docx"))

    if not docx_files:
        print("\n❌ 当前目录下没有找到 .docx 文件！")
        print(f"   请将 .docx 文件放到: {WORK_DIR}")
        sys.exit(1)

    print(f"\n找到 {len(docx_files)} 个 DOCX 文件:")
    for f in docx_files:
        print(f"  - {f.name}")

    # 创建输出目录
    OUTPUT_DIR.mkdir(exist_ok=True)

    # 逐个处理
    success = 0
    for filepath in docx_files:
        try:
            if process_docx(filepath):
                success += 1
        except Exception as e:
            print(f"  ❌ 处理失败: {e}")

    print(f"\n{'='*60}")
    print(f"🎉 完成! 成功处理 {success}/{len(docx_files)} 个文件")
    print(f"📁 JSON 文件输出目录: {OUTPUT_DIR}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
